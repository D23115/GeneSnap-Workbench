"""Version-consistent editable exports for the SYN workflow."""

from __future__ import annotations

import hashlib
import json
import re
import shutil
from dataclasses import asdict, dataclass
from datetime import datetime
from decimal import Decimal
from enum import Enum
from pathlib import Path

from Bio import SeqIO
from Bio.Seq import Seq
from Bio.SeqFeature import FeatureLocation, SeqFeature
from Bio.SeqRecord import SeqRecord
from docx import Document
from docx.oxml.ns import qn
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill

from genesnap_workbench.domain.syn import (
    OligoMixPlan,
    OligoResuspensionPlan,
    SYNColonyPCRRecord,
    SYNDesignVersion,
)


class SYNExportError(ValueError):
    """Raised before formal files are exposed when export data is incomplete."""


@dataclass(frozen=True, slots=True)
class GeneratedArtifact:
    artifact_type: str
    design_version_id: str
    generated_at: datetime
    path: Path
    content_sha256: str
    lifecycle_status: str = "current"


@dataclass(frozen=True, slots=True)
class SYNExportBundle:
    batch_dir: Path
    artifacts: tuple[GeneratedArtifact, ...]

    def path_for(self, artifact_type: str) -> Path:
        for artifact in self.artifacts:
            if artifact.artifact_type == artifact_type:
                return artifact.path
        raise KeyError(artifact_type)


def _json_default(value):
    if isinstance(value, Decimal):
        return str(value)
    if isinstance(value, datetime):
        return value.isoformat()
    if isinstance(value, Enum):
        return value.value
    raise TypeError(f"Unsupported JSON value: {type(value).__name__}")


def _safe_name(value: str) -> str:
    cleaned = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", value.strip())
    cleaned = cleaned.rstrip(". ")
    return cleaned or "SYN-target"


def _sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _validate_export_inputs(
    design: SYNDesignVersion,
    resuspension_plan: OligoResuspensionPlan | None,
    mix_plan: OligoMixPlan | None,
) -> None:
    if design.qc_result.blocked_reasons:
        reasons = "；".join(design.qc_result.blocked_reasons)
        raise SYNExportError(f"设计仍有阻断项，不能正式导出：{reasons}")
    confirmed = any(
        record.field_path == "design_confirmation"
        and record.new_value == "confirmed"
        for record in design.manual_overrides
    )
    if design.requires_confirmation and not confirmed:
        raise SYNExportError("设计含需人工确认项，确认并记录原因后才能正式导出")
    for label, plan in (("复溶", resuspension_plan), ("混池", mix_plan)):
        if plan is not None and plan.design_version_id != design.design_version_id:
            raise SYNExportError(f"{label}计划与当前设计版本不一致")
    if mix_plan is not None and not mix_plan.is_formal_export_ready:
        raise SYNExportError("混池计划缺少公共加样体积或实际加样体积")


def _write_metadata(sheet, design: SYNDesignVersion, generated_at: datetime) -> None:
    for row, (label, value) in enumerate(
        (
            ("设计版本", design.design_version_id),
            ("目标序列校验值", design.final_checksum),
            ("生成时间", generated_at.isoformat()),
        ),
        start=1,
    ):
        sheet.cell(row, 1, label).font = Font(bold=True)
        sheet.cell(row, 2, value)


def _write_header(sheet, headers: tuple[str, ...], row: int = 5) -> None:
    for column, value in enumerate(headers, start=1):
        cell = sheet.cell(row, column, value)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="2F6F62")
    sheet.freeze_panes = f"A{row + 1}"
    sheet.auto_filter.ref = f"A{row}:{sheet.cell(row, len(headers)).coordinate}"


def _write_order_workbook(
    design: SYNDesignVersion,
    generated_at: datetime,
    path: Path,
) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "oligo 订购"
    _write_metadata(sheet, design, generated_at)
    _write_header(
        sheet,
        ("oligo 名称", "序列（5'-3'）", "pool", "module", "方向", "长度 nt"),
    )
    for row, oligo in enumerate(design.oligos, start=6):
        for column, value in enumerate(
            (
                oligo.name,
                oligo.sequence,
                oligo.pool_id,
                oligo.module_id,
                oligo.strand,
                len(oligo.sequence),
            ),
            start=1,
        ):
            sheet.cell(row, column, value)
    sheet.column_dimensions["A"].width = 24
    sheet.column_dimensions["B"].width = 72
    for column in ("C", "D", "E", "F"):
        sheet.column_dimensions[column].width = 16
    workbook.save(path)


def _write_resuspension_workbook(
    design: SYNDesignVersion,
    plan: OligoResuspensionPlan,
    generated_at: datetime,
    path: Path,
) -> None:
    names = {oligo.oligo_id: oligo.name for oligo in design.oligos}
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "oligo 复溶"
    _write_metadata(sheet, design, generated_at)
    _write_header(
        sheet,
        (
            "oligo 名称",
            "计划量 nmol",
            "实际量 nmol",
            "加水体积 uL",
            "目标浓度 uM",
            "实际浓度 uM",
        ),
    )
    for row, item in enumerate(plan.items, start=6):
        values = (
            names.get(item.oligo_id, item.oligo_id),
            item.planned_amount_nmol,
            item.actual_amount_nmol,
            item.water_volume_ul,
            item.target_stock_concentration_uM,
            item.actual_stock_concentration_uM,
        )
        for column, value in enumerate(values, start=1):
            sheet.cell(
                row,
                column,
                float(value) if isinstance(value, Decimal) else value,
            )
    workbook.save(path)


def _write_mix_workbook(
    design: SYNDesignVersion,
    plan: OligoMixPlan,
    generated_at: datetime,
    path: Path,
) -> None:
    names = {oligo.oligo_id: oligo.name for oligo in design.oligos}
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "oligo 混池"
    _write_metadata(sheet, design, generated_at)
    _write_header(
        sheet,
        ("pool", "oligo 名称", "参考浓度 uM", "实际浓度 uM", "加样体积 uL"),
    )
    for row, item in enumerate(plan.items, start=6):
        values = (
            item.pool_id,
            names.get(item.oligo_id, item.oligo_id),
            item.reference_concentration_uM,
            item.actual_concentration_uM,
            item.sample_volume_ul,
        )
        for column, value in enumerate(values, start=1):
            sheet.cell(
                row,
                column,
                float(value) if isinstance(value, Decimal) else value,
            )
    workbook.save(path)


def _write_design_report(
    design: SYNDesignVersion,
    target_name: str,
    generated_at: datetime,
    path: Path,
) -> None:
    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    document.add_heading("GeneSnap SYN 设计报告", level=0)
    document.add_paragraph(f"目标名称：{target_name}")
    document.add_paragraph(f"设计版本：{design.design_version_id}")
    document.add_paragraph(f"目标序列校验值：{design.final_checksum}")
    document.add_paragraph(
        f"预期质粒校验值：{design.plasmid_simulation.expected_plasmid_checksum}",
    )
    document.add_paragraph(f"生成时间：{generated_at.isoformat()}")
    document.add_heading("Assembly oligo", level=1)
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ("名称", "序列（5'-3'）", "pool", "module", "Tm (C)")
    for cell, value in zip(table.rows[0].cells, headers, strict=True):
        cell.text = value
    for oligo in design.oligos:
        cells = table.add_row().cells
        values = (
            oligo.name,
            oligo.sequence,
            oligo.pool_id,
            oligo.module_id,
            str(oligo.tm_metadata.tm_celsius),
        )
        for cell, value in zip(cells, values, strict=True):
            cell.text = value
    document.add_heading("模块与组装路线", level=1)
    document.add_paragraph(
        f"路线：{design.module_plan.route.value}；"
        f"原因：{design.module_plan.routing_reason}",
    )
    for module in design.module_plan.modules:
        document.add_paragraph(
            f"{module.module_id}：{module.start + 1}-{module.end} bp；"
            f"边界说明：{module.boundary_reason}",
        )
    document.add_heading("序列 QC", level=1)
    if not design.qc_result.risks:
        document.add_paragraph("未发现当前规则定义的风险项。")
    for risk in design.qc_result.risks:
        document.add_paragraph(
            f"[{risk.severity}] {risk.start + 1}-{risk.end}：{risk.message}",
        )
    document.add_paragraph(
        "本报告表示序列满足当前软件规则，不代表湿实验必然成功。",
    )
    document.save(path)


def _write_colony_workbook(
    design: SYNDesignVersion,
    colonies: tuple[SYNColonyPCRRecord, ...],
    generated_at: datetime,
    path: Path,
) -> None:
    latest = {}
    for colony in colonies:
        latest[colony.clone_id] = colony
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "菌落 PCR"
    _write_metadata(sheet, design, generated_at)
    _write_header(
        sheet,
        ("克隆名称", "轮次记录", "预期条带 bp", "结果", "备注"),
    )
    ordered = sorted(
        latest.values(),
        key=lambda item: (item.attempt_id, item.clone_no),
    )
    for row, colony in enumerate(ordered, start=6):
        for column, value in enumerate(
            (
                colony.display_name,
                colony.attempt_id,
                colony.expected_band_bp,
                colony.result.value,
                colony.observed_note,
            ),
            start=1,
        ):
            sheet.cell(row, column, value)
    workbook.save(path)


def _write_genbank(
    design: SYNDesignVersion,
    target_name: str,
    path: Path,
) -> None:
    simulation = design.plasmid_simulation
    record = SeqRecord(
        Seq(simulation.expected_plasmid_sequence),
        id=_safe_name(target_name)[:16],
        name=_safe_name(target_name)[:16],
        description=f"GeneSnap SYN expected plasmid {design.design_version_id}",
    )
    record.annotations["molecule_type"] = "DNA"
    record.annotations["topology"] = "circular"
    record.annotations["structured_comment"] = {
        "GeneSnap-Data": {
            "Design Version": design.design_version_id,
            "Final Checksum": design.final_checksum,
            "Plasmid Checksum": simulation.expected_plasmid_checksum,
            "Vector Record": simulation.vector_record_id,
            "Protocol Version": simulation.protocol_version_id,
        },
    }
    for feature in simulation.features:
        record.features.append(
            SeqFeature(
                FeatureLocation(feature.start, feature.end, strand=feature.strand),
                type=feature.feature_type,
                qualifiers={"label": [feature.label]},
            ),
        )
    SeqIO.write(record, path, "genbank")


def _unique_batch_dir(
    output_root: Path,
    target_name: str,
    design_version_id: str,
    generated_at: datetime,
) -> Path:
    timestamp = generated_at.strftime("%Y%m%dT%H%M%S")
    base = f"{_safe_name(target_name)}_{_safe_name(design_version_id)}_{timestamp}"
    candidate = output_root / base
    suffix = 2
    while candidate.exists() or candidate.with_name(f".{candidate.name}.tmp").exists():
        candidate = output_root / f"{base}_{suffix}"
        suffix += 1
    return candidate


def export_syn_bundle(
    design: SYNDesignVersion,
    output_root: Path,
    *,
    target_name: str,
    generated_at: datetime,
    resuspension_plan: OligoResuspensionPlan | None = None,
    mix_plan: OligoMixPlan | None = None,
    colonies: tuple[SYNColonyPCRRecord, ...] = (),
) -> SYNExportBundle:
    """Generate one batch and expose it only after every file succeeds."""
    _validate_export_inputs(design, resuspension_plan, mix_plan)
    root = Path(output_root)
    root.mkdir(parents=True, exist_ok=True)
    batch_dir = _unique_batch_dir(
        root,
        target_name,
        design.design_version_id,
        generated_at,
    )
    staging_dir = batch_dir.with_name(f".{batch_dir.name}.tmp")
    staging_dir.mkdir()
    files: list[tuple[str, str]] = []
    try:
        json_path = staging_dir / "01_设计数据.json"
        json_path.write_text(
            json.dumps(
                asdict(design),
                ensure_ascii=False,
                indent=2,
                sort_keys=True,
                default=_json_default,
            )
            + "\n",
            encoding="utf-8",
        )
        files.append(("design_json", json_path.name))

        order_path = staging_dir / "02_oligo订购表.xlsx"
        _write_order_workbook(design, generated_at, order_path)
        files.append(("oligo_order_xlsx", order_path.name))

        if resuspension_plan is not None:
            resuspension_path = staging_dir / "03_oligo复溶表.xlsx"
            _write_resuspension_workbook(
                design,
                resuspension_plan,
                generated_at,
                resuspension_path,
            )
            files.append(("oligo_resuspension_xlsx", resuspension_path.name))

        if mix_plan is not None:
            mix_path = staging_dir / "04_oligo混池表.xlsx"
            _write_mix_workbook(design, mix_plan, generated_at, mix_path)
            files.append(("oligo_mix_xlsx", mix_path.name))

        report_path = staging_dir / "05_SYN设计报告.docx"
        _write_design_report(design, target_name, generated_at, report_path)
        files.append(("design_report_docx", report_path.name))

        colony_path = staging_dir / "06_菌落PCR记录表.xlsx"
        _write_colony_workbook(design, colonies, generated_at, colony_path)
        files.append(("colony_pcr_xlsx", colony_path.name))

        genbank_path = staging_dir / "07_预期重组质粒.gb"
        _write_genbank(design, target_name, genbank_path)
        files.append(("expected_plasmid_genbank", genbank_path.name))

        staging_dir.rename(batch_dir)
    except Exception:
        shutil.rmtree(staging_dir, ignore_errors=True)
        raise

    artifacts = tuple(
        GeneratedArtifact(
            artifact_type=artifact_type,
            design_version_id=design.design_version_id,
            generated_at=generated_at,
            path=batch_dir / filename,
            content_sha256=_sha256_file(batch_dir / filename),
        )
        for artifact_type, filename in files
    )
    return SYNExportBundle(batch_dir=batch_dir, artifacts=artifacts)
