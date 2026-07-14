"""表达类设计、订购、送测和预期质粒的版本一致输出。"""

from __future__ import annotations

import hashlib
import json
import re
from dataclasses import asdict, dataclass
from datetime import date, datetime
from enum import Enum
from pathlib import Path

from Bio import SeqIO
from Bio.Seq import Seq
from Bio.SeqFeature import FeatureLocation, SeqFeature
from Bio.SeqRecord import SeqRecord
from docx import Document
from docx.oxml.ns import qn
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill

from genesnap_workbench.domain.expression import (
    ExpressionCloneResultRecord,
    ExpressionDesignVersion,
    ExpressionVectorDesignResult,
)
from genesnap_workbench.template_engine.workbook_templates import (
    ContactProfile,
    LocalWorkbookTemplateStore,
)
from genesnap_workbench.project_workflow.project_folders import ProjectWorkspace

from .syn_exports import GeneratedArtifact


class ExpressionExportError(ValueError):
    pass


@dataclass(frozen=True, slots=True)
class ExpressionExportBundle:
    artifacts: tuple[GeneratedArtifact, ...]

    def path_for(self, artifact_type: str) -> Path:
        matches = self.paths_for(artifact_type)
        if len(matches) != 1:
            raise KeyError(f"{artifact_type} has {len(matches)} paths")
        return matches[0]

    def paths_for(self, artifact_type: str) -> tuple[Path, ...]:
        return tuple(
            item.path for item in self.artifacts if item.artifact_type == artifact_type
        )


def _json_default(value):
    if isinstance(value, datetime):
        return value.isoformat()
    if isinstance(value, Enum):
        return value.value
    raise TypeError(f"Unsupported JSON value: {type(value).__name__}")


def _safe_name(value: str) -> str:
    cleaned = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", value.strip()).rstrip(". ")
    return cleaned or "未命名"


def _unique_path(path: Path) -> Path:
    if not path.exists():
        return path
    number = 2
    while True:
        candidate = path.with_name(f"{path.stem}_{number}{path.suffix}")
        if not candidate.exists():
            return candidate
        number += 1


def _sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _write_metadata(
    sheet,
    design: ExpressionDesignVersion,
    vector_result: ExpressionVectorDesignResult,
    generated_at: datetime,
) -> None:
    rows = (
        ("设计版本", design.design_version_id),
        ("目标基因", design.gene_symbol),
        ("载体序列校验值", vector_result.vector_checksum),
        ("生成时间", generated_at.isoformat()),
    )
    for row, (label, value) in enumerate(rows, start=1):
        sheet.cell(row, 1, label).font = Font(bold=True)
        sheet.cell(row, 2, value)


def _write_header(sheet, headers: tuple[str, ...]) -> None:
    for column, value in enumerate(headers, start=1):
        cell = sheet.cell(6, column, value)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="2F6F62")
    sheet.freeze_panes = "A7"
    sheet.auto_filter.ref = f"A6:{sheet.cell(6, len(headers)).coordinate}"


def _construct_pairs(
    design: ExpressionDesignVersion,
    vector_result: ExpressionVectorDesignResult,
):
    plans = {item.construct_id: item for item in vector_result.construct_plans}
    if set(plans) != {item.construct_id for item in design.constructs}:
        raise ExpressionExportError("表达设计与载体结果中的构建不一致")
    return tuple((construct, plans[construct.construct_id]) for construct in design.constructs)


def _write_primer_order(
    design: ExpressionDesignVersion,
    vector_result: ExpressionVectorDesignResult,
    generated_at: datetime,
    path: Path,
) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "引物订购"
    _write_metadata(sheet, design, vector_result, generated_at)
    _write_header(
        sheet,
        ("引物名称", "引物序列(5'-3')", "目标基因", "构建", "方向", "长度 nt"),
    )
    row = 7
    for construct, plan in _construct_pairs(design, vector_result):
        for primer in plan.primers:
            values = (
                primer.name,
                primer.sequence,
                design.gene_symbol,
                construct.construct_name,
                primer.direction,
                len(primer.sequence),
            )
            for column, value in enumerate(values, start=1):
                sheet.cell(row, column, value)
            row += 1
    sheet.column_dimensions["A"].width = 32
    sheet.column_dimensions["B"].width = 72
    workbook.save(path)


def _primer_records(
    design: ExpressionDesignVersion,
    vector_result: ExpressionVectorDesignResult,
) -> tuple[dict[str, object], ...]:
    records: list[dict[str, object]] = []
    for construct, plan in _construct_pairs(design, vector_result):
        for primer in plan.primers:
            records.append(
                {
                    "primer_name": primer.name,
                    "sequence": primer.sequence,
                    "gene_symbol": design.gene_symbol,
                    "direction": primer.direction,
                    "length": len(primer.sequence),
                    "note": construct.construct_name,
                },
            )
    return tuple(records)


def _write_sequencing_order(
    design: ExpressionDesignVersion,
    vector_result: ExpressionVectorDesignResult,
    generated_at: datetime,
    path: Path,
    *,
    clones_per_construct: int,
    sequencing_method: str,
) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "送测"
    _write_metadata(sheet, design, vector_result, generated_at)
    _write_header(
        sheet,
        ("样本名称", "测序方式", "目标基因", "构建", "克隆号", "备注"),
    )
    row = 7
    for construct in design.constructs:
        for clone_no in range(1, clones_per_construct + 1):
            values = (
                f"{construct.construct_name}-{clone_no}",
                sequencing_method,
                design.gene_symbol,
                construct.construct_name,
                clone_no,
                "整质粒测通" if sequencing_method.lower() == "nanopore" else "",
            )
            for column, value in enumerate(values, start=1):
                sheet.cell(row, column, value)
            row += 1
    sheet.column_dimensions["A"].width = 34
    for column in ("B", "C", "D", "E", "F"):
        sheet.column_dimensions[column].width = 18
    workbook.save(path)


def _sequencing_records(
    design: ExpressionDesignVersion,
    *,
    clones_per_construct: int,
    sequencing_method: str,
) -> tuple[dict[str, object], ...]:
    return tuple(
        {
            "sample_name": f"{construct.construct_name}-{clone_no}",
            "gene_symbol": design.gene_symbol,
            "clone_no": clone_no,
            "method": sequencing_method,
            "note": "整质粒测通" if sequencing_method.lower() == "nanopore" else "",
        }
        for construct in design.constructs
        for clone_no in range(1, clones_per_construct + 1)
    )


def _write_report(
    design: ExpressionDesignVersion,
    vector_result: ExpressionVectorDesignResult,
    generated_at: datetime,
    path: Path,
) -> None:
    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    document.add_heading("GeneSnap 表达类质粒构建设计报告", level=0)
    document.add_paragraph(f"项目号：{design.project_id}")
    document.add_paragraph(f"目标基因：{design.gene_symbol}")
    document.add_paragraph(f"物种：{design.species}")
    document.add_paragraph(f"设计版本：{design.design_version_id}")
    document.add_paragraph(f"载体序列校验值：{vector_result.vector_checksum}")
    document.add_paragraph(f"生成时间：{generated_at.isoformat()}")
    table = document.add_table(rows=1, cols=8)
    table.style = "Table Grid"
    headers = (
        "构建",
        "类型",
        "编码区 bp",
        "插入片段 bp",
        "PCR片段数",
        "正向引物",
        "反向引物",
        "标签/融合",
    )
    for cell, value in zip(table.rows[0].cells, headers, strict=True):
        cell.text = value
    for construct, plan in _construct_pairs(design, vector_result):
        cells = table.add_row().cells
        values = (
            construct.construct_name,
            construct.kind.value,
            str(len(construct.coding_sequence)),
            str(len(construct.insert_sequence)),
            str(len(construct.fragments)),
            "\n".join(
                f"{item.name}: {item.sequence}"
                for item in plan.primers
                if item.direction == "F"
            ),
            "\n".join(
                f"{item.name}: {item.sequence}"
                for item in plan.primers
                if item.direction == "R"
            ),
            construct.c_terminal_fusion_name or "无",
        )
        for cell, value in zip(cells, values, strict=True):
            cell.text = value
    if design.design_warnings:
        document.add_heading("人工复核项", level=1)
        for warning in design.design_warnings:
            document.add_paragraph(warning, style="List Bullet")
    document.add_paragraph("本报告表示序列满足当前软件规则，不代表湿实验必然成功。")
    document.save(path)


def _write_expected_genbank(
    design: ExpressionDesignVersion,
    construct_index: int,
    plan,
    path: Path,
) -> None:
    construct = design.constructs[construct_index]
    record = SeqRecord(
        Seq(plan.expected_plasmid_sequence),
        id=f"{_safe_name(design.gene_symbol)}-{_safe_name(construct.construct_name)}"[:16],
        name=f"{_safe_name(design.gene_symbol)}-{_safe_name(construct.construct_name)}"[:16],
        description=f"GeneSnap expected expression plasmid {design.design_version_id}",
    )
    record.annotations["molecule_type"] = "DNA"
    record.annotations["topology"] = "circular"
    record.annotations["structured_comment"] = {
        "GeneSnap-Data": {
            "Design Version": design.design_version_id,
            "Expected Plasmid Checksum": plan.expected_plasmid_checksum,
            "Protocol Version": design.protocol_version_id,
        },
    }
    insert_start = plan.expected_plasmid_sequence.find(construct.insert_sequence)
    if insert_start >= 0:
        record.features.append(
            SeqFeature(
                FeatureLocation(
                    insert_start,
                    insert_start + len(construct.insert_sequence),
                    strand=1,
                ),
                type="CDS",
                qualifiers={"label": [f"{design.gene_symbol}-{construct.construct_name}"]},
            ),
        )
    SeqIO.write(record, path, "genbank")


def export_expression_bundle(
    design: ExpressionDesignVersion,
    vector_result: ExpressionVectorDesignResult,
    workspace: ProjectWorkspace,
    *,
    generated_at: datetime,
    primer_order_date: date,
    sequencing_order_date: date,
    clones_per_construct: int = 5,
    primer_vendor_name: str = "标准",
    sequencing_vendor_name: str = "标准",
    sequencing_method: str = "Nanopore",
    workbook_template_store: LocalWorkbookTemplateStore | None = None,
    contact_profile: ContactProfile = ContactProfile(),
    primer_template_id: str | None = None,
    sequencing_template_id: str | None = None,
) -> ExpressionExportBundle:
    if design.requires_confirmation:
        raise ExpressionExportError("设计仍有未确认项，不能生成正式订购文件")
    if vector_result.design_version_id != design.design_version_id:
        raise ExpressionExportError("表达设计与载体结果版本不一致")
    if not 1 <= clones_per_construct <= 96:
        raise ExpressionExportError("每个构建的送测克隆数必须在 1 到 96 之间")
    pairs = _construct_pairs(design, vector_result)

    design_dir = workspace.folder("01_design")
    order_dir = workspace.folder("02_orders")
    report_dir = workspace.folder("04_reports")
    files: list[tuple[str, Path]] = []

    json_path = _unique_path(
        design_dir / f"{design.gene_symbol}-{design.design_version_id}-设计数据.json",
    )
    json_path.write_text(
        json.dumps(
            {"design": asdict(design), "vector_result": asdict(vector_result)},
            ensure_ascii=False,
            indent=2,
            sort_keys=True,
            default=_json_default,
        )
        + "\n",
        encoding="utf-8",
    )
    files.append(("design_json", json_path))

    primer_path = _unique_path(
        order_dir
        / f"{primer_order_date:%Y%m%d}-{_safe_name(primer_vendor_name)}引物订购表-{_safe_name(design.gene_symbol)}.xlsx",
    )
    if primer_template_id:
        if workbook_template_store is None:
            raise ExpressionExportError("已选择引物模板，但模板库不可用")
        workbook_template_store.render(
            primer_template_id,
            records=_primer_records(design, vector_result),
            contact=contact_profile,
            output_path=primer_path,
        )
    else:
        _write_primer_order(design, vector_result, generated_at, primer_path)
    files.append(("primer_order_xlsx", primer_path))

    sequencing_path = _unique_path(
        order_dir
        / f"{sequencing_order_date:%Y%m%d}-{_safe_name(sequencing_vendor_name)}测序表-{_safe_name(design.gene_symbol)}.xlsx",
    )
    if sequencing_template_id:
        if workbook_template_store is None:
            raise ExpressionExportError("已选择送测模板，但模板库不可用")
        workbook_template_store.render(
            sequencing_template_id,
            records=_sequencing_records(
                design,
                clones_per_construct=clones_per_construct,
                sequencing_method=sequencing_method,
            ),
            contact=contact_profile,
            output_path=sequencing_path,
        )
    else:
        _write_sequencing_order(
            design,
            vector_result,
            generated_at,
            sequencing_path,
            clones_per_construct=clones_per_construct,
            sequencing_method=sequencing_method,
        )
    files.append(("sequencing_order_xlsx", sequencing_path))

    report_path = _unique_path(
        report_dir / f"表达类质粒构建设计报告-{_safe_name(design.gene_symbol)}.docx",
    )
    _write_report(design, vector_result, generated_at, report_path)
    files.append(("design_report_docx", report_path))

    for index, (construct, plan) in enumerate(pairs):
        map_path = _unique_path(
            design_dir
            / f"预期质粒-{_safe_name(construct.construct_name)}.gb",
        )
        _write_expected_genbank(design, index, plan, map_path)
        files.append(("expected_plasmid_genbank", map_path))

    return ExpressionExportBundle(
        artifacts=tuple(
            GeneratedArtifact(
                artifact_type=artifact_type,
                design_version_id=design.design_version_id,
                generated_at=generated_at,
                path=path,
                content_sha256=_sha256_file(path),
            )
            for artifact_type, path in files
        ),
    )


def export_expression_analysis_report(
    design: ExpressionDesignVersion,
    vector_result: ExpressionVectorDesignResult,
    records: tuple[ExpressionCloneResultRecord, ...],
    output_dir: Path,
    *,
    analyzed_at: datetime,
    unmatched_files: tuple[Path, ...] = (),
    ambiguous_files: tuple[Path, ...] = (),
) -> GeneratedArtifact:
    output_root = Path(output_dir)
    output_root.mkdir(parents=True, exist_ok=True)
    path = _unique_path(
        output_root
        / f"{analyzed_at:%Y%m%d}-表达类测序分析-{_safe_name(design.gene_symbol)}.xlsx",
    )
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "克隆判读"
    _write_metadata(sheet, design, vector_result, analyzed_at)
    _write_header(
        sheet,
        (
            "克隆名称",
            "构建",
            "结果",
            "判读说明",
            "覆盖率",
            "一致性",
            "替换",
            "插入 bp",
            "缺失 bp",
            "移码",
            "提前终止",
            "测序文件",
        ),
    )
    construct_names = {
        item.construct_id: item.construct_name for item in design.constructs
    }
    for row, record in enumerate(records, start=7):
        values = (
            record.clone_name,
            construct_names.get(record.construct_id, ""),
            record.status.upper(),
            record.reason,
            record.coverage,
            record.identity,
            record.substitution_count,
            record.insertion_count,
            record.deletion_count,
            "是" if record.frameshift else "否",
            "是" if record.premature_stop else "否",
            "\n".join(record.source_files),
        )
        for column, value in enumerate(values, start=1):
            sheet.cell(row, column, value)
    sheet.column_dimensions["A"].width = 30
    sheet.column_dimensions["D"].width = 58
    sheet.column_dimensions["L"].width = 72
    if unmatched_files or ambiguous_files:
        issues = workbook.create_sheet("文件匹配问题")
        issues.append(("类型", "文件"))
        for path_item in unmatched_files:
            issues.append(("未匹配文件", str(path_item)))
        for path_item in ambiguous_files:
            issues.append(("匹配到多个克隆", str(path_item)))
        issues.column_dimensions["A"].width = 22
        issues.column_dimensions["B"].width = 90
    workbook.save(path)
    return GeneratedArtifact(
        artifact_type="sequencing_analysis_xlsx",
        design_version_id=design.design_version_id,
        generated_at=analyzed_at,
        path=path,
        content_sha256=_sha256_file(path),
    )
