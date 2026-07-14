"""shRNA 设计、订购、送测和预期质粒的版本一致输出。"""

from __future__ import annotations

import hashlib
import json
import re
from dataclasses import asdict, dataclass
from datetime import date, datetime
from decimal import Decimal
from enum import Enum
from pathlib import Path

from Bio.Seq import Seq
from Bio.SeqFeature import FeatureLocation, SeqFeature
from Bio.SeqRecord import SeqRecord
from docx import Document
from docx.oxml.ns import qn
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill

from genesnap_workbench.domain.shrna import ShRNACloneResultRecord, ShRNADesignVersion
from genesnap_workbench.template_engine.workbook_templates import (
    ContactProfile,
    LocalWorkbookTemplateStore,
)
from genesnap_workbench.project_workflow.project_folders import ProjectWorkspace

from .genbank_io import write_genbank_utf8
from .syn_exports import GeneratedArtifact


class ShRNAExportError(ValueError):
    pass


@dataclass(frozen=True, slots=True)
class ShRNAExportBundle:
    artifacts: tuple[GeneratedArtifact, ...]

    def path_for(self, artifact_type: str) -> Path:
        matches = self.paths_for(artifact_type)
        if len(matches) != 1:
            raise KeyError(f"{artifact_type} has {len(matches)} paths")
        return matches[0]

    def paths_for(self, artifact_type: str) -> tuple[Path, ...]:
        return tuple(
            item.path for item in self.artifacts if item.artifact_type == artifact_type
        )


def _json_default(value):
    if isinstance(value, Decimal):
        return str(value)
    if isinstance(value, datetime):
        return value.isoformat()
    if isinstance(value, Enum):
        return value.value
    raise TypeError(f"Unsupported JSON value: {type(value).__name__}")


def _safe_name(value: str) -> str:
    cleaned = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", value.strip()).rstrip(". ")
    return cleaned or "未命名"


def _unique_path(path: Path) -> Path:
    if not path.exists():
        return path
    number = 2
    while True:
        candidate = path.with_name(f"{path.stem}_{number}{path.suffix}")
        if not candidate.exists():
            return candidate
        number += 1


def _sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _write_metadata(sheet, design: ShRNADesignVersion, generated_at: datetime) -> None:
    rows = (
        ("设计版本", design.design_version_id),
        ("目标基因", design.gene_symbol),
        ("载体序列校验值", design.vector_checksum),
        ("生成时间", generated_at.isoformat()),
    )
    for row, (label, value) in enumerate(rows, start=1):
        sheet.cell(row, 1, label).font = Font(bold=True)
        sheet.cell(row, 2, value)


def _write_header(sheet, headers: tuple[str, ...]) -> None:
    for column, value in enumerate(headers, start=1):
        cell = sheet.cell(6, column, value)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="2F6F62")
    sheet.freeze_panes = "A7"
    sheet.auto_filter.ref = f"A6:{sheet.cell(6, len(headers)).coordinate}"


def _write_primer_order(
    design: ShRNADesignVersion,
    generated_at: datetime,
    path: Path,
) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "引物订购"
    _write_metadata(sheet, design, generated_at)
    _write_header(
        sheet,
        ("引物名称", "引物序列(5'-3')", "目标基因", "Target", "方向", "长度 nt"),
    )
    row = 7
    for target in design.targets:
        for direction, name, sequence in (
            ("F", target.oligos.forward_name, target.oligos.forward_sequence),
            ("R", target.oligos.reverse_name, target.oligos.reverse_sequence),
        ):
            values = (
                name,
                sequence,
                design.gene_symbol,
                target.target_no,
                direction,
                len(sequence),
            )
            for column, value in enumerate(values, start=1):
                sheet.cell(row, column, value)
            row += 1
    sheet.column_dimensions["A"].width = 24
    sheet.column_dimensions["B"].width = 72
    workbook.save(path)


def _primer_records(design: ShRNADesignVersion) -> tuple[dict[str, object], ...]:
    records: list[dict[str, object]] = []
    for target in design.targets:
        for direction, name, sequence in (
            ("F", target.oligos.forward_name, target.oligos.forward_sequence),
            ("R", target.oligos.reverse_name, target.oligos.reverse_sequence),
        ):
            records.append(
                {
                    "primer_name": name,
                    "sequence": sequence,
                    "gene_symbol": design.gene_symbol,
                    "gene_id": design.gene_id,
                    "transcript_accession": design.transcript_accession,
                    "direction": direction,
                    "length": len(sequence),
                    "product_length": len(sequence),
                },
            )
    return tuple(records)


def _write_sequencing_order(
    design: ShRNADesignVersion,
    generated_at: datetime,
    path: Path,
) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "送测"
    _write_metadata(sheet, design, generated_at)
    _write_header(
        sheet,
        ("样本名称", "测序引物", "目标基因", "Target", "克隆号", "测序方式"),
    )
    row = 7
    for target in design.targets:
        for clone_no, clone_name in enumerate(target.clone_names, start=1):
            values = (
                clone_name,
                "U6",
                design.gene_symbol,
                target.target_no,
                clone_no,
                "Sanger",
            )
            for column, value in enumerate(values, start=1):
                sheet.cell(row, column, value)
            row += 1
    sheet.column_dimensions["A"].width = 28
    for column in ("B", "C", "D", "E", "F"):
        sheet.column_dimensions[column].width = 16
    workbook.save(path)


def _sequencing_records(design: ShRNADesignVersion) -> tuple[dict[str, object], ...]:
    records: list[dict[str, object]] = []
    for target in design.targets:
        for clone_no, clone_name in enumerate(target.clone_names, start=1):
            records.append(
                {
                    "sample_name": clone_name,
                    "primer_name": "U6",
                    "gene_symbol": design.gene_symbol,
                    "clone_no": clone_no,
                    "method": "Sanger",
                },
            )
    return tuple(records)


def _write_report(
    design: ShRNADesignVersion,
    generated_at: datetime,
    path: Path,
) -> None:
    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    document.add_heading("GeneSnap shRNA 设计报告", level=0)
    document.add_paragraph(f"项目号：{design.project_id}")
    document.add_paragraph(f"目标基因：{design.gene_symbol}")
    document.add_paragraph(f"物种：{design.species}")
    document.add_paragraph(f"设计版本：{design.design_version_id}")
    document.add_paragraph(f"载体序列校验值：{design.vector_checksum}")
    document.add_paragraph(f"生成时间：{generated_at.isoformat()}")
    table = document.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    headers = ("Target", "靶序列", "位置", "得分", "BLAST", "正向 oligo", "反向 oligo")
    for cell, value in zip(table.rows[0].cells, headers, strict=True):
        cell.text = value
    for target in design.targets:
        cells = table.add_row().cells
        values = (
            str(target.target_no),
            target.candidate.target_sequence,
            "" if target.candidate.start_position is None else str(target.candidate.start_position),
            str(target.candidate.intrinsic_score),
            target.candidate.blast_status.value,
            target.oligos.forward_sequence,
            target.oligos.reverse_sequence,
        )
        for cell, value in zip(cells, values, strict=True):
            cell.text = value
    document.add_paragraph(
        "本报告表示序列满足当前软件规则，不代表湿实验必然成功。",
    )
    document.save(path)


def _write_expected_genbank(
    design: ShRNADesignVersion,
    target_index: int,
    path: Path,
) -> None:
    target = design.targets[target_index]
    simulation = design.plasmid_simulations[target_index]
    record = SeqRecord(
        Seq(simulation.expected_plasmid_sequence),
        id=f"{_safe_name(design.gene_symbol)}-sh{target.target_no}"[:16],
        name=f"{_safe_name(design.gene_symbol)}-sh{target.target_no}"[:16],
        description=f"GeneSnap expected shRNA plasmid {design.design_version_id}",
    )
    record.annotations["molecule_type"] = "DNA"
    record.annotations["topology"] = "circular"
    record.annotations["structured_comment"] = {
        "GeneSnap-Data": {
            "Design Version": design.design_version_id,
            "Vector Checksum": design.vector_checksum,
            "Plasmid Checksum": simulation.expected_plasmid_checksum,
            "Protocol Version": design.vector_protocol_version_id,
        },
    }
    insert_start = simulation.left_cut_position
    insert_end = insert_start + len(target.oligos.forward_sequence)
    record.features.append(
        SeqFeature(
            FeatureLocation(insert_start, insert_end, strand=1),
            type="misc_feature",
            qualifiers={"label": [f"{design.gene_symbol}-shRNA-{target.target_no}"]},
        ),
    )
    write_genbank_utf8(record, path)


def export_shrna_bundle(
    design: ShRNADesignVersion,
    workspace: ProjectWorkspace,
    *,
    generated_at: datetime,
    primer_order_date: date,
    sequencing_order_date: date,
    primer_vendor_name: str = "标准",
    sequencing_vendor_name: str = "标准",
    workbook_template_store: LocalWorkbookTemplateStore | None = None,
    contact_profile: ContactProfile = ContactProfile(),
    primer_template_id: str | None = None,
    sequencing_template_id: str | None = None,
) -> ShRNAExportBundle:
    if design.requires_confirmation:
        raise ShRNAExportError("设计仍有未确认的 BLAST 项，不能生成正式订购文件")

    design_dir = workspace.folder("01_design")
    order_dir = workspace.folder("02_orders")
    report_dir = workspace.folder("04_reports")
    files: list[tuple[str, Path]] = []

    json_path = _unique_path(design_dir / f"{design.gene_symbol}-{design.design_version_id}-设计数据.json")
    json_path.write_text(
        json.dumps(
            asdict(design),
            ensure_ascii=False,
            indent=2,
            sort_keys=True,
            default=_json_default,
        )
        + "\n",
        encoding="utf-8",
    )
    files.append(("design_json", json_path))

    primer_path = _unique_path(
        order_dir
        / f"{primer_order_date:%Y%m%d}-{_safe_name(primer_vendor_name)}引物订购表-{_safe_name(design.gene_symbol)}.xlsx",
    )
    if primer_template_id:
        if workbook_template_store is None:
            raise ShRNAExportError("已选择引物模板，但模板库不可用")
        workbook_template_store.render(
            primer_template_id,
            records=_primer_records(design),
            contact=contact_profile,
            document_values={"order_date": generated_at.date()},
            output_path=primer_path,
        )
    else:
        _write_primer_order(design, generated_at, primer_path)
    files.append(("primer_order_xlsx", primer_path))

    sequencing_path = _unique_path(
        order_dir
        / f"{sequencing_order_date:%Y%m%d}-{_safe_name(sequencing_vendor_name)}测序表-{_safe_name(design.gene_symbol)}.xlsx",
    )
    if sequencing_template_id:
        if workbook_template_store is None:
            raise ShRNAExportError("已选择送测模板，但模板库不可用")
        workbook_template_store.render(
            sequencing_template_id,
            records=_sequencing_records(design),
            contact=contact_profile,
            output_path=sequencing_path,
        )
    else:
        _write_sequencing_order(design, generated_at, sequencing_path)
    files.append(("sequencing_order_xlsx", sequencing_path))

    report_path = _unique_path(
        report_dir / f"shRNA质粒构建设计报告-{_safe_name(design.gene_symbol)}.docx",
    )
    _write_report(design, generated_at, report_path)
    files.append(("design_report_docx", report_path))

    for index, target in enumerate(design.targets):
        map_path = _unique_path(
            design_dir / f"预期质粒-{_safe_name(design.gene_symbol)}-shRNA-{target.target_no}.gb",
        )
        _write_expected_genbank(design, index, map_path)
        files.append(("expected_plasmid_genbank", map_path))

    return ShRNAExportBundle(
        artifacts=tuple(
            GeneratedArtifact(
                artifact_type=artifact_type,
                design_version_id=design.design_version_id,
                generated_at=generated_at,
                path=path,
                content_sha256=_sha256_file(path),
            )
            for artifact_type, path in files
        ),
    )


def export_shrna_analysis_report(
    design: ShRNADesignVersion,
    records: tuple[ShRNACloneResultRecord, ...],
    output_dir: Path,
    *,
    analyzed_at: datetime,
    unmatched_files: tuple[Path, ...] = (),
    ambiguous_files: tuple[Path, ...] = (),
) -> GeneratedArtifact:
    output_root = Path(output_dir)
    output_root.mkdir(parents=True, exist_ok=True)
    path = _unique_path(
        output_root
        / f"{analyzed_at:%Y%m%d}-shRNA测序分析-{_safe_name(design.gene_symbol)}.xlsx",
    )
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "克隆判读"
    _write_metadata(sheet, design, analyzed_at)
    _write_header(
        sheet,
        ("克隆名称", "Target", "结果", "判读说明", "匹配位置", "测序文件"),
    )
    target_numbers = {
        target.target_id: target.target_no for target in design.targets
    }
    for row, record in enumerate(records, start=7):
        values = (
            record.clone_name,
            target_numbers.get(record.target_id, ""),
            record.status.upper(),
            record.reason,
            record.match_start,
            "\n".join(record.source_files),
        )
        for column, value in enumerate(values, start=1):
            sheet.cell(row, column, value)
    sheet.column_dimensions["A"].width = 26
    sheet.column_dimensions["D"].width = 52
    sheet.column_dimensions["F"].width = 72
    if unmatched_files or ambiguous_files:
        issues = workbook.create_sheet("文件匹配问题")
        issues.append(("类型", "文件"))
        for path_item in unmatched_files:
            issues.append(("未匹配文件", str(path_item)))
        for path_item in ambiguous_files:
            issues.append(("匹配到多个克隆", str(path_item)))
        issues.column_dimensions["A"].width = 22
        issues.column_dimensions["B"].width = 90
    workbook.save(path)
    return GeneratedArtifact(
        artifact_type="sequencing_analysis_xlsx",
        design_version_id=design.design_version_id,
        generated_at=analyzed_at,
        path=path,
        content_sha256=_sha256_file(path),
    )
