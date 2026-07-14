import json
import tempfile
import unittest
from dataclasses import replace
from datetime import datetime, timezone
from decimal import Decimal
from pathlib import Path

from Bio import SeqIO
from docx import Document
from openpyxl import load_workbook

from genesnap_workbench.domain.syn import (
    OligoMixItem,
    OligoMixPlan,
    OligoResuspensionItem,
    OligoResuspensionPlan,
    ResuspensionStatus,
)
from genesnap_workbench.sequence_core.syn_design import confirm_syn_design_warnings
from genesnap_workbench.template_engine.syn_exports import (
    SYNExportError,
    export_syn_bundle,
)
from genesnap_workbench.vector_library.syn import circular_sequence_checksum
from tests.test_syn_models import make_design_version


NOW = datetime(2026, 7, 12, 11, 0, tzinfo=timezone.utc)


def make_export_design():
    design = make_design_version()
    plasmid = replace(
        design.plasmid_simulation,
        expected_plasmid_checksum=circular_sequence_checksum(
            design.plasmid_simulation.expected_plasmid_sequence,
        ),
    )
    return replace(design, plasmid_simulation=plasmid)


def make_resuspension_plan():
    return OligoResuspensionPlan(
        design_version_id="design-v1",
        items=(
            OligoResuspensionItem(
                oligo_id="oligo-1",
                planned_amount_nmol=Decimal("25"),
                actual_amount_nmol=Decimal("27.4"),
                target_stock_concentration_uM=Decimal("100"),
                water_volume_ul=Decimal("274"),
                actual_stock_concentration_uM=Decimal("100"),
            ),
            OligoResuspensionItem(
                oligo_id="oligo-2",
                planned_amount_nmol=Decimal("25"),
                actual_amount_nmol=Decimal("24.8"),
                target_stock_concentration_uM=Decimal("100"),
                water_volume_ul=Decimal("248"),
                actual_stock_concentration_uM=Decimal("100"),
            ),
        ),
        status=ResuspensionStatus.COMPLETE,
        generated_at=NOW,
    )


def make_mix_plan(standard_volume=Decimal("1.5")):
    return OligoMixPlan(
        design_version_id="design-v1",
        standard_volume_per_oligo_ul=standard_volume,
        items=(
            OligoMixItem(
                pool_id="pool-1",
                oligo_id="oligo-1",
                reference_concentration_uM=Decimal("100"),
                actual_concentration_uM=Decimal("100"),
                sample_volume_ul=standard_volume,
            ),
            OligoMixItem(
                pool_id="pool-1",
                oligo_id="oligo-2",
                reference_concentration_uM=Decimal("100"),
                actual_concentration_uM=Decimal("100"),
                sample_volume_ul=standard_volume,
            ),
        ),
        generated_at=NOW,
    )


class SYNExportTests(unittest.TestCase):
    def test_confirmation_required_design_must_have_audited_confirmation(self):
        design = replace(
            make_export_design(),
            design_warnings=("局部结构需要人工复核",),
            requires_confirmation=True,
        )
        with tempfile.TemporaryDirectory() as temp_dir:
            with self.assertRaisesRegex(SYNExportError, "人工确认"):
                export_syn_bundle(
                    design,
                    Path(temp_dir),
                    target_name="SYN-target",
                    generated_at=NOW,
                )
            confirmed = confirm_syn_design_warnings(
                design,
                override_id="override-confirm-design",
                reason="已人工复核 overlap 与模块边界",
                actor="tester",
                occurred_at=NOW,
            )
            bundle = export_syn_bundle(
                confirmed,
                Path(temp_dir),
                target_name="SYN-target",
                generated_at=NOW,
            )
            self.assertTrue(bundle.path_for("design_json").exists())
            self.assertEqual(
                confirmed.manual_overrides[-1].field_path,
                "design_confirmation",
            )

    def test_bundle_files_share_design_version_and_checksum(self):
        design = make_export_design()
        with tempfile.TemporaryDirectory() as temp_dir:
            bundle = export_syn_bundle(
                design,
                Path(temp_dir),
                target_name="SYN-target",
                generated_at=NOW,
                resuspension_plan=make_resuspension_plan(),
                mix_plan=make_mix_plan(),
            )

            artifact_types = {item.artifact_type for item in bundle.artifacts}
            self.assertEqual(
                artifact_types,
                {
                    "design_json",
                    "oligo_order_xlsx",
                    "oligo_resuspension_xlsx",
                    "oligo_mix_xlsx",
                    "design_report_docx",
                    "colony_pcr_xlsx",
                    "expected_plasmid_genbank",
                },
            )
            for artifact in bundle.artifacts:
                self.assertEqual(artifact.design_version_id, "design-v1")
                self.assertTrue(artifact.path.exists())
                self.assertEqual(len(artifact.content_sha256), 64)

            payload = json.loads(bundle.path_for("design_json").read_text("utf-8"))
            self.assertEqual(payload["design_version_id"], "design-v1")
            self.assertEqual(payload["final_checksum"], design.final_checksum)

            report_text = "\n".join(
                paragraph.text
                for paragraph in Document(
                    bundle.path_for("design_report_docx"),
                ).paragraphs
            )
            self.assertIn("design-v1", report_text)
            self.assertIn(design.final_checksum, report_text)

            record = SeqIO.read(
                bundle.path_for("expected_plasmid_genbank"),
                "genbank",
            )
            self.assertEqual(
                circular_sequence_checksum(str(record.seq)),
                design.plasmid_simulation.expected_plasmid_checksum,
            )
            metadata = record.annotations["structured_comment"]["GeneSnap-Data"]
            self.assertEqual(metadata["Design Version"], "design-v1")
            self.assertEqual(metadata["Final Checksum"], design.final_checksum)

    def test_oligo_rows_are_consistent_across_xlsx_docx_and_json(self):
        design = make_export_design()
        with tempfile.TemporaryDirectory() as temp_dir:
            bundle = export_syn_bundle(
                design,
                Path(temp_dir),
                target_name="SYN-target",
                generated_at=NOW,
            )
            workbook = load_workbook(
                bundle.path_for("oligo_order_xlsx"),
                data_only=True,
            )
            sheet = workbook["oligo 订购"]
            xlsx_rows = tuple(
                (sheet.cell(row, 1).value, sheet.cell(row, 2).value)
                for row in range(6, 6 + len(design.oligos))
            )
            expected = tuple((item.name, item.sequence) for item in design.oligos)
            self.assertEqual(xlsx_rows, expected)
            self.assertEqual(sheet["B1"].value, design.design_version_id)
            self.assertEqual(sheet["B2"].value, design.final_checksum)

            payload = json.loads(bundle.path_for("design_json").read_text("utf-8"))
            json_rows = tuple(
                (item["name"], item["sequence"]) for item in payload["oligos"]
            )
            self.assertEqual(json_rows, expected)

            document = Document(bundle.path_for("design_report_docx"))
            table_rows = tuple(
                (row.cells[0].text, row.cells[1].text)
                for row in document.tables[0].rows[1:]
            )
            self.assertEqual(table_rows, expected)

    def test_resuspension_uses_actual_amount_and_unready_mix_is_rejected(self):
        design = make_export_design()
        invalid_mix = make_mix_plan(standard_volume=None)
        with tempfile.TemporaryDirectory() as temp_dir:
            with self.assertRaisesRegex(SYNExportError, "混池"):
                export_syn_bundle(
                    design,
                    Path(temp_dir),
                    target_name="SYN-target",
                    generated_at=NOW,
                    resuspension_plan=make_resuspension_plan(),
                    mix_plan=invalid_mix,
                )
            self.assertEqual(tuple(Path(temp_dir).iterdir()), ())

            bundle = export_syn_bundle(
                design,
                Path(temp_dir),
                target_name="SYN-target",
                generated_at=NOW,
                resuspension_plan=make_resuspension_plan(),
            )
            workbook = load_workbook(
                bundle.path_for("oligo_resuspension_xlsx"),
                data_only=True,
            )
            sheet = workbook["oligo 复溶"]
            self.assertEqual(sheet["C6"].value, 27.4)
            self.assertEqual(sheet["D6"].value, 274)

    def test_reexport_creates_new_auditable_directory_without_overwrite(self):
        design = make_export_design()
        with tempfile.TemporaryDirectory() as temp_dir:
            first = export_syn_bundle(
                design,
                Path(temp_dir),
                target_name="SYN-target",
                generated_at=NOW,
            )
            second = export_syn_bundle(
                design,
                Path(temp_dir),
                target_name="SYN-target",
                generated_at=NOW,
            )

            self.assertNotEqual(first.batch_dir, second.batch_dir)
            self.assertTrue(first.path_for("design_json").exists())
            self.assertTrue(second.path_for("design_json").exists())
            self.assertEqual(
                first.path_for("design_json").read_bytes(),
                second.path_for("design_json").read_bytes(),
            )


if __name__ == "__main__":
    unittest.main()
